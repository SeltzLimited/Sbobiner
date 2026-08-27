"""Esportazione: testo semplice, sottotitoli SRT/VTT, Word DOCX formattato."""
from docx import Document
from docx.shared import Pt


def _hms(sec, sep=":", ms=False):
    h, r = divmod(int(sec), 3600)
    m, s = divmod(r, 60)
    if ms:
        return f"{h:02d}:{m:02d}:{s:02d}{sep}{int((sec % 1) * 1000):03d}"
    return f"{h:02d}:{m:02d}:{s:02d}"


def _spk(seg):
    return f"{seg['speaker']}: " if seg.get("speaker") else ""


def to_txt(result):
    lines = []
    for sec in result["sections"]:
        lines.append(f"\n== Sezione {sec['index']} ({_hms(sec['start'])}) ==")
        for seg in sec["segments"]:
            lines.append(f"[{_hms(seg['start'])}] {_spk(seg)}{seg['text']}")
    return "\n".join(lines).strip() + "\n"


def _subtitles(result, vtt):
    blocks = ["WEBVTT\n"] if vtt else []
    sep = "." if vtt else ","
    for i, seg in enumerate(result["segments"], 1):
        start, end = _hms(seg["start"], sep, ms=True), _hms(seg["end"], sep, ms=True)
        head = "" if vtt else f"{i}\n"
        blocks.append(f"{head}{start} --> {end}\n{_spk(seg)}{seg['text']}\n")
    return "\n".join(blocks).strip() + "\n"


def to_srt(result):
    return _subtitles(result, vtt=False)


def to_vtt(result):
    return _subtitles(result, vtt=True)


def to_docx(result, path, source_name="", mode="lezione"):
    doc = Document()
    doc.add_heading(f"Trascrizione — {source_name}", level=0)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Lingua: {result['language']}  ·  "
        f"Durata: {_hms(result['segments'][-1]['end']) if result['segments'] else '00:00:00'}  ·  "
        f"Sezioni: {len(result['sections'])}  ·  "
        f"Modalità: {mode}"
    ).italic = True

    if mode == "riunione" and (result["decisions"] or result["action_items"]):
        _decisions_actions(doc, result)

    doc.add_heading("Trascrizione", level=1)
    for sec in result["sections"]:
        doc.add_heading(f"Sezione {sec['index']}  ({_hms(sec['start'])})", level=2)
        for seg in sec["segments"]:
            p = doc.add_paragraph()
            if seg.get("speaker"):
                p.add_run(f"{seg['speaker']} ").bold = True
            p.add_run(f"[{_hms(seg['start'])}] ").font.size = Pt(8)
            p.add_run(seg["text"])

    if mode == "lezione" and (result["decisions"] or result["action_items"]):
        _decisions_actions(doc, result)

    doc.save(str(path))
    return path


def _decisions_actions(doc, result):
    if result["decisions"]:
        doc.add_heading("Decisioni", level=1)
        for d in result["decisions"]:
            doc.add_paragraph(f"[{_hms(d['start'])}] {d['text']}", style="List Bullet")
    if result["action_items"]:
        doc.add_heading("Action item", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for h, cell in zip(("Testo", "Speaker", "Quando"), table.rows[0].cells):
            cell.paragraphs[0].add_run(h).bold = True
        for a in result["action_items"]:
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text = a["text"], a["speaker"], _hms(a["start"])


WRITERS = {"txt": to_txt, "srt": to_srt, "vtt": to_vtt}


def write(result, fmt, path, source_name="", mode="lezione"):
    if fmt == "docx":
        return to_docx(result, path, source_name, mode)
    with open(path, "w", encoding="utf-8") as f:
        f.write(WRITERS[fmt](result))
    return path
